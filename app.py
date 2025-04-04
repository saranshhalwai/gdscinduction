import os
import streamlit as st
from groq import Groq
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument
from docx import Document as DocxDocument
import whisper
from TTS.api import TTS
from streamlit_mic_recorder import mic_recorder  # Streamlit mic recorder
from dotenv import load_dotenv
load_dotenv()

st.set_page_config(page_title="Chatbot with RAG + Mic", layout="centered")

# Set API key
api_key = os.getenv("GROQ_API_KEY")
if not api_key:
    st.error("GROQ_API_KEY is not set! Please configure your environment variable.")
    st.stop()

# Initialize Groq Client
client = Groq(api_key=api_key)

# Load and process the document
file_path = "ES103.docx"
try:
    doc = DocxDocument(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    if not text.strip():
        raise ValueError("The document is empty!")
    docs = [LangchainDocument(page_content=text)]
except Exception as e:
    st.error(f"Error loading file: {e}")
    st.stop()

# Split document into chunks
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
all_splits = text_splitter.split_documents(docs)

# Convert to embeddings and store in FAISS
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
db = FAISS.from_documents(all_splits, embeddings)
db.save_local("faiss_index")

# Initialize session state for conversation history
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# Function for multi-step retrieval
def multi_step_retrieval(query):
    retrieved_docs = db.similarity_search(query, k=3)
    context = "\n".join([doc.page_content for doc in retrieved_docs])
    
    # If the context is too short, try expanding the search
    if len(context) < 500:
        additional_query = " ".join(query.split()[:5])  # Use first few words as a refined query
        additional_docs = db.similarity_search(additional_query, k=3)
        additional_context = "\n".join([doc.page_content for doc in additional_docs])
        context += "\n" + additional_context
    
    return context

# Function to handle conversation-style chatbot response
def chat_with_rag(query):
    context = multi_step_retrieval(query)
    
    messages = [{"role": "system", "content": "You are a helpful chatbot using retrieved knowledge."}]
    messages += st.session_state.chat_history  # Include past messages
    messages.append({"role": "user", "content": f"User Query: {query}\n\nRelevant Context:\n{context}\n\nChatbot:"})
    
    completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=messages,
        temperature=0.6,
        max_tokens=500,
        top_p=0.95,
        stream=False,
    )
    
    response = completion.choices[0].message.content
    
    # Store conversation history
    st.session_state.chat_history.append({"role": "user", "content": query})
    st.session_state.chat_history.append({"role": "assistant", "content": response})
    
    return response

# Initialize Whisper for STT
whisper_model = whisper.load_model("base")

def transcribe_audio(audio_data):
    with open("temp_audio.wav", "wb") as f:
        f.write(audio_data)
    result = whisper_model.transcribe("temp_audio.wav")
    return result["text"]

# Initialize TTS
tts = TTS("tts_models/en/ljspeech/glow-tts")

def text_to_speech(text):
    tts.tts_to_file(text=text, file_path="output.wav")

# Streamlit UI
st.title("Chatbot with RAG + Built-in Mic STT")

# Text Input
user_input = st.text_input("Ask something:")
if user_input:
    response = chat_with_rag(user_input)
    st.write(response)
    text_to_speech(response)
    st.audio("output.wav")

# Microphone Input
st.write("**Record your voice to ask a question:**")
audio = mic_recorder()
if audio and "bytes" in audio:
    st.write("Processing audio...")
    transcribed_text = transcribe_audio(audio["bytes"])
    st.write("Transcribed Text:", transcribed_text)
    response = chat_with_rag(transcribed_text)
    st.write(response)
    text_to_speech(response)
    st.audio("output.wav")