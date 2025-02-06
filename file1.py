import os
import streamlit as st
from groq import Groq
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.docstore.document import Document
from docx import Document as DocxDocument
import whisper
from TTS.api import TTS
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Set API key
api_key = os.getenv("GROQ_API_KEY")
if not api_key:
    st.error("GROQ_API_KEY is not set! Please configure your environment variable.")
    st.stop()

# Initialize Groq Client
client = Groq(api_key=api_key)

# Load and process the document
file_path = "ES103.docx"

try:
    doc = DocxDocument(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    if not text.strip():
        raise ValueError("The document is empty!")
    
    # Split the document into smaller chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,  # chunk size (characters)
        chunk_overlap=200,  # chunk overlap (characters)
        add_start_index=True,  # track index in original document
    )
    docs = [Document(page_content=text)]
    all_splits = text_splitter.split_documents(docs)
except Exception as e:
    st.error(f"Error loading file: {e}")
    st.stop()

# Convert to embeddings and store in FAISS
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
db = FAISS.from_documents(all_splits, embeddings)

# Save FAISS index
db.save_local("faiss_index")

def chat_with_rag(query):
    # Retrieve relevant documents (Reduced k=2 to avoid long context)
    retrieved_docs = db.similarity_search(query, k=3)
    context = "\n".join([doc.page_content for doc in retrieved_docs])
    
    # Generate response using Groq API
    completion = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": f"User Query: {query}\n\nRelevant Context:\n{context}\n\nChatbot:"}],
        temperature=0.6,
        max_tokens=500,
        top_p=0.95,
        stream=False,
    )
    return completion.choices[0].message.content

# Initialize Whisper for speech-to-text
whisper_model = whisper.load_model("base")

def transcribe_audio(audio_path):
    result = whisper_model.transcribe(audio_path)
    return result["text"]

# Initialize TTS
tts = TTS("tts_models/en/ljspeech/glow-tts")

def text_to_speech(text):
    tts.tts_to_file(text=text, file_path="output.wav")

# Streamlit UI
st.title("Chatbot with RAG + Voice")

# Text Input
user_input = st.text_input("Ask something:")
if user_input:
    response = chat_with_rag(user_input)
    st.write(response)
    
    # Convert response to speech
    text_to_speech(response)
    st.audio("output.wav")

# Voice Input (File Upload)
uploaded_audio = st.file_uploader("Upload a voice file:", type=["wav", "mp3"])
if uploaded_audio:
    with open("temp_audio.wav", "wb") as f:
        f.write(uploaded_audio.read())
    
    transcribed_text = transcribe_audio("temp_audio.wav")
    st.write("Transcribed Text:", transcribed_text)
    
    # Get response from chatbot
    response = chat_with_rag(transcribed_text)
    st.write(response)
    
    # Convert to speech
    text_to_speech(response)
    st.audio("output.wav")
